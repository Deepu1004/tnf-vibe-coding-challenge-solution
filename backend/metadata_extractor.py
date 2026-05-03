import hashlib
import os
from typing import Dict, Any, Optional
import mimetypes

try:
    import fitz  # PyMuPDF for PDFs
except ImportError:
    fitz = None

try:
    from docx import Document  # python-docx for Word files
except ImportError:
    Document = None


def _hash_file(file_path: str) -> str:
    """Generate a stable SHA-256 hash for the full file content."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as file_handle:
        for chunk in iter(lambda: file_handle.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from PDF file"""
    if not fitz:
        return {"error": "PyMuPDF not installed"}
    
    try:
        doc = fitz.open(file_path)
        metadata = doc.metadata or {}
        
        # Extract standard PDF metadata
        result = {
            "author": metadata.get("author", ""),
            "title": metadata.get("title", ""),
            "subject": metadata.get("subject", ""),
            "creator": metadata.get("creator", ""),
            "producer": metadata.get("producer", ""),
            "creation_date": str(metadata.get("creationDate", "")),
            "mod_date": str(metadata.get("modDate", "")),
            "pages": len(doc),
            "file_size": os.path.getsize(file_path),
            "content_hash": _hash_file(file_path),
        }
        
        # Extract text from first 2 pages for style analysis
        text_sample = ""
        for page_num in range(min(2, len(doc))):
            text_sample += doc[page_num].get_text() + "\n"
        
        result["text_sample"] = text_sample[:2000]  # First 2000 chars
        result["total_words"] = len(text_sample.split())
        
        # Add classification
        classification = classify_document(text_sample)
        result.update(classification)
        
        doc.close()
        return result
    except Exception as e:
        return {"error": f"Failed to extract PDF metadata: {str(e)}"}


def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from DOCX file"""
    if not Document:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(file_path)
        props = doc.core_properties
        
        result = {
            "author": props.author or "",
            "title": props.title or "",
            "subject": props.subject or "",
            "creator": "Microsoft Word",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "file_size": os.path.getsize(file_path),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "content_hash": _hash_file(file_path),
        }
        
        # Extract text sample
        text_sample = ""
        for para in doc.paragraphs[:20]:  # First 20 paragraphs
            text_sample += para.text + "\n"
        
        result["text_sample"] = text_sample[:2000]
        result["total_words"] = len(text_sample.split())
        
        # Add classification
        classification = classify_document(text_sample)
        result.update(classification)
        
        return result
    except Exception as e:
        return {"error": f"Failed to extract DOCX metadata: {str(e)}"}


def extract_txt_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        result = {
            "author": "",
            "file_size": os.path.getsize(file_path),
            "text_sample": content[:2000],
            "total_words": len(content.split()),
            "total_lines": len(content.split('\n')),
            "content_hash": _hash_file(file_path),
        }
        
        # Add classification
        classification = classify_document(content[:2000])
        result.update(classification)
        
        return result
    except Exception as e:
        return {"error": f"Failed to extract TXT metadata: {str(e)}"}


def extract_file_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from uploaded file based on file type
    """
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Route to appropriate extractor
    if ext == ".pdf":
        return extract_pdf_metadata(file_path)
    elif ext == ".docx":
        return extract_docx_metadata(file_path)
    elif ext == ".txt":
        return extract_txt_metadata(file_path)
    else:
        return {
            "error": f"Unsupported file type: {ext}",
            "file_size": os.path.getsize(file_path),
            "file_type": ext,
            "content_hash": _hash_file(file_path),
        }

def classify_document(text: str) -> Dict[str, str]:
    text_lower = text.lower()
    
    doc_type = "paper"
    if "chapter" in text_lower or "book" in text_lower:
        doc_type = "book"
    elif "journal" in text_lower or "volume" in text_lower or "issue" in text_lower:
        doc_type = "journal"
        
    doc_kind = "science"
    if any(k in text_lower for k in ["engineering", "architecture", "design", "mechanic"]):
        doc_kind = "engineering"
    elif any(k in text_lower for k in ["medical", "health", "clinical", "patient", "disease"]):
        doc_kind = "medical"
        
    return {
        "document_type": doc_type,
        "document_kind": doc_kind
    }



def analyze_document_similarity(
    metadata1: Dict[str, Any],
    metadata2: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Analyze similarity between two documents
    """
    similarity_score = 0
    matching_fields = []
    
    # Author match
    if (metadata1.get("author") and metadata2.get("author") and
        metadata1.get("author") == metadata2.get("author")):
        similarity_score += 25
        matching_fields.append("author")
    
    # Creator/Software match
    if (metadata1.get("creator") and metadata2.get("creator") and
        metadata1.get("creator") == metadata2.get("creator")):
        similarity_score += 15
        matching_fields.append("creator")
    
    # Title similarity
    if (metadata1.get("title") and metadata2.get("title") and
        metadata1.get("title").lower() == metadata2.get("title").lower()):
        similarity_score += 20
        matching_fields.append("title")
    
    # Text similarity (basic word overlap)
    text1 = metadata1.get("text_sample", "").lower()
    text2 = metadata2.get("text_sample", "").lower()
    
    if text1 and text2:
        words1 = set(text1.split())
        words2 = set(text2.split())
        if len(words1) > 0 and len(words2) > 0:
            overlap = len(words1 & words2) / len(words1 | words2)
            text_similarity = int(overlap * 100)
            similarity_score += text_similarity // 4  # Weight it less
            if overlap > 0.3:
                matching_fields.append("text_content")
    
    return {
        "similarity_score": min(similarity_score, 100),
        "matching_fields": matching_fields
    }
